# -*- coding: utf-8 -*-
"""将 notebook_cell_sources 中的代码写入 word_gen_system_demo.ipynb。运行: python _patch_nb.py"""
import importlib.util
import json
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent
NB = ROOT / "word_gen_system_demo.ipynb"


def L(s: str) -> list:
    if not s.endswith("\n"):
        s += "\n"
    return [ln if ln.endswith("\n") else ln + "\n" for ln in s.splitlines(keepends=True)]


def set_cell_by_prefix(cells, prefix: str, new_source: str) -> bool:
    for c in cells:
        if c.get("cell_type") != "code":
            continue
        src = "".join(c.get("source", []))
        if src.strip().startswith(prefix):
            c["source"] = L(new_source)
            c["outputs"] = []
            c["execution_count"] = None
            return True
    print('WARN: cell not found:', prefix, file=sys.stderr)
    return False


def main():
    spec = importlib.util.spec_from_file_location(
        "notebook_cell_sources", ROOT / "notebook_cell_sources.py"
    )
    ncs = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(ncs)

    nb = json.loads(NB.read_text(encoding="utf-8"))
    cells = nb["cells"]

    for c in cells:
        if c.get("cell_type") != "code":
            continue
        src = "".join(c.get("source", []))
        if "'product_name': '产品1'" in src:
            src = src.replace(
                "    {'branch_name': 'A支行', 'product_name': '产品1', 'overdue_rate': 1.2},\n"
                "    {'branch_name': 'A支行', 'product_name': '产品2', 'overdue_rate': 0.4},\n",
                "    {'branch_name': 'A支行', 'product_name': '法人e抵', 'overdue_rate': 1.2},\n"
                "    {'branch_name': 'A支行', 'product_name': '个人e抵押', 'overdue_rate': 0.4},\n",
            )
            c["source"] = L(src)
            break

    CELL5 = r'''# Cell 5: 生成示例 Word 模板（docxtpl 表格行循环 + 条件 + AI 段标记）
from docx import Document
from docx.shared import Pt

template_path = DIRS['templates'] / 'template_1.docx'

AI_BLOCK_MARKERS = ['§AIBLOCK0§']


def build_demo_template(path: Path) -> None:
    """
    构建演示用 Word 模板。
    表格循环：三行结构 {%tr for %} / 数据行 / {%tr endfor %}。
    """
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(11)

    doc.add_heading('普惠支行报告', level=1)
    p = doc.add_paragraph()
    p.add_run('支行：')
    p.add_run('{{ data.branch_name }}')
    p.add_run('，报告期：')
    p.add_run('{{ data.report_date }}')

    p2 = doc.add_paragraph()
    p2.add_run('普惠贷款余额为 ')
    p2.add_run('{{ loan_summary.loan_balance }}')
    p2.add_run(' 万元；房产评估金额 ')
    p2.add_run('{{ loan_summary.house_eval_amount }}')
    p2.add_run(' 万元。')

    doc.add_paragraph(
        '清偿能力判断：{% if loan_summary.house_eval_amount < loan_summary.loan_balance %}无法清偿{% else %}可清偿{% endif %}。'
    )

    doc.add_paragraph('各产品不良情况：')
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    table.cell(0, 0).text = '产品名称'
    table.cell(0, 1).text = '不良率(%)'
    table.cell(1, 0).text = '{%tr for p in products %}'
    table.cell(1, 1).text = ''
    table.cell(2, 0).text = '{{ p.product_name }}'
    table.cell(2, 1).text = '{{ p.overdue_rate }}'
    table.cell(3, 0).text = '{%tr endfor %}'
    table.cell(3, 1).text = ''

    doc.add_paragraph('风险分析（以下段落将由 AI 改写）：')
    ai_para = doc.add_paragraph(
        '§AIBLOCK0§基于不良环比{{ data.bad_mom }}、同比{{ data.bad_yoy }}，请撰写风险分析与建议。'
    )
    comment_text = (
        'prompt="基于不良额环比{{ data.bad_mom }}、同比{{ data.bad_yoy }}数据，'
        '分析该支行风险状况，给出针对性建议"'
    )
    try:
        doc.add_comment(runs=ai_para.runs, text=comment_text, author='AI', initials='AI')
    except Exception as e:
        raise RuntimeError('python-docx 需 >=1.2 以支持 add_comment') from e

    doc.save(path)


build_demo_template(template_path)
print('Wrote template:', template_path)
'''
    set_cell_by_prefix(cells, "# Cell 5:", CELL5)
    set_cell_by_prefix(cells, "# Cell 8:", ncs.CELL8)
    set_cell_by_prefix(cells, "# Cell 10:", ncs.CELL10)
    set_cell_by_prefix(cells, "# Cell 11:", ncs.CELL11)
    set_cell_by_prefix(cells, "# Cell 12:", ncs.CELL12)
    set_cell_by_prefix(cells, "# Cell 13:", ncs.CELL13)

    # 最后一格：交互入口说明（优先使用末尾的空 code cell）
    for c in reversed(cells):
        if c.get("cell_type") == "code" and not "".join(c.get("source", [])).strip():
            c["source"] = L(
                "# Cell 14: 运行智能生成（交互）\n"
                "# 执行下面一行后按提示输入模板 ID（逗号分隔）、是否 AI、并确认数据表已就绪。\n"
                "# interactive_smart_generation()\n"
                "\n"
                "# 或一行非交互测试（不调用 AI、自动确认）：\n"
                "# run_smart_document_generation([1], use_ai=False, skip_confirm=True)\n"
            )
            break

    NB.write_text(json.dumps(nb, ensure_ascii=False, indent=2), encoding="utf-8")
    print("OK:", NB)


if __name__ == "__main__":
    main()
