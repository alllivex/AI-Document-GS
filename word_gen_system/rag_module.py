"""
增强版RAG知识库模块，支持Word文档解析和灵活知识库引用语法
"""
import os
import json
import re
from pathlib import Path
from typing import List, Dict, Optional, Tuple, Any, Union
from sentence_transformers import SentenceTransformer
import numpy as np
from PyPDF2 import PdfReader
from docx import Document
import yaml

try:
    from sklearn.metrics.pairwise import cosine_similarity
    HAS_SKLEARN = True
except ImportError:
    HAS_SKLEARN = False
    print("警告：未安装scikit-learn，将使用numpy计算相似度")


class KBManager:
    """知识库管理器"""
    
    def __init__(self, kb_base_dir: Path | str):
        self.kb_base_dir = Path(kb_base_dir)
        self.use_vector_search = True
        
        # 尝试加载向量模型
        try:
            self.model = SentenceTransformer('all-MiniLM-L6-v2')
            # 如果是中文环境，使用更好的中文模型
            try:
                # 尝试加载中文模型
                self.model = SentenceTransformer('paraphrase-multilingual-MiniLM-L12-v2')
            except:
                # 如果加载失败，使用默认模型
                self.model = SentenceTransformer('all-MiniLM-L6-v2')
        except Exception as e:
            print(f"警告：向量模型加载失败，将使用关键词匹配模式: {e}")
            self.model = None
            self.use_vector_search = False
            
    def _extract_text_from_pdf(self, pdf_path: Path) -> str:
        """从PDF文件中提取文本"""
        reader = PdfReader(pdf_path)
        text = ""
        for page in reader.pages:
            text += page.extract_text()
        return text
        
    def _extract_text_from_txt(self, txt_path: Path) -> str:
        """从TXT文件中提取文本"""
        with open(txt_path, 'r', encoding='utf-8') as f:
            return f.read()
            
    def _extract_text_from_yaml(self, yaml_path: Path) -> str:
        """从YAML文件中提取文本"""
        with open(yaml_path, 'r', encoding='utf-8') as f:
            data = yaml.safe_load(f)
            return json.dumps(data, ensure_ascii=False)
    
    def _extract_text_from_docx(self, docx_path: Path) -> str:
        """从Word文档中提取文本"""
        try:
            doc = Document(docx_path)
            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)
            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text.append(cell.text)
            return '\n'.join(full_text)
        except Exception as e:
            print(f"警告：读取Word文档 {docx_path} 失败: {e}")
            return ""
    
    def _get_document_content(self, doc_path: Path) -> str:
        """获取文档内容，根据文件类型选择相应方法"""
        ext = doc_path.suffix.lower()
        if ext == '.pdf':
            return self._extract_text_from_pdf(doc_path)
        elif ext == '.txt':
            return self._extract_text_from_txt(doc_path)
        elif ext in ['.yaml', '.yml']:
            return self._extract_text_from_yaml(doc_path)
        elif ext == '.docx':
            return self._extract_text_from_docx(doc_path)
        elif ext == '.doc':
            # 尝试作为docx处理，如果失败则返回空
            return self._extract_text_from_docx(doc_path)
        else:
            # 默认按文本文件处理
            try:
                with open(doc_path, 'r', encoding='utf-8') as f:
                    return f.read()
            except UnicodeDecodeError:
                # 如果UTF-8解码失败，尝试其他编码
                with open(doc_path, 'r', encoding='gbk') as f:
                    return f.read()
    
    def _chunk_text(self, text: str, chunk_size: int = 512) -> List[str]:
        """将长文本切分成块"""
        chunks = []
        for i in range(0, len(text), chunk_size):
            chunk = text[i:i + chunk_size]
            chunks.append(chunk)
        return chunks
    
    def _embed_chunks(self, chunks: List[str]) -> np.ndarray:
        """对文本块进行向量化"""
        embeddings = self.model.encode(chunks)
        return np.array(embeddings)
    
    def _calculate_similarity(self, query_embedding: np.ndarray, doc_embeddings: np.ndarray) -> np.ndarray:
        """计算查询与文档块的相似度"""
        # 使用余弦相似度
        query_norm = query_embedding / np.linalg.norm(query_embedding)
        doc_norm = doc_embeddings / np.linalg.norm(doc_embeddings, axis=1, keepdims=True)
        similarity = np.dot(doc_norm, query_norm.T)
        return similarity.flatten()
    
    def load_kb(self, kb_name: str) -> Dict:
        """加载知识库"""
        kb_path = self.kb_base_dir / kb_name
        if not kb_path.exists():
            raise ValueError(f"知识库 {kb_name} 不存在")
        
        # 加载元数据
        meta_file = kb_path / "metadata.json"
        if not meta_file.exists():
            raise ValueError(f"知识库 {kb_name} 缺少 metadata.json 文件")
        
        with open(meta_file, 'r', encoding='utf-8') as f:
            metadata = json.load(f)
        
        # 加载文档并分块
        documents = metadata.get("documents", [])
        all_chunks = []
        all_embeddings = []
        
        for doc_name in documents:
            doc_path = kb_path / doc_name
            if not doc_path.exists():
                print(f"警告: 知识库 {kb_name} 中文档 {doc_name} 不存在")
                continue
            
            content = self._get_document_content(doc_path)
            chunks = self._chunk_text(content)
            
            # 为每个块创建信息
            for i, chunk in enumerate(chunks):
                all_chunks.append({
                    "content": chunk,
                    "doc_name": doc_name,
                    "chunk_idx": i,
                    "kb_name": kb_name
                })
        
        # 批量嵌入
        if all_chunks:
            texts = [chunk["content"] for chunk in all_chunks]
            embeddings = self._embed_chunks(texts)
        else:
            embeddings = np.array([])
        
        return {
            "metadata": metadata,
            "chunks": all_chunks,
            "embeddings": embeddings
        }
    
    def search_kb(self, kb_name: str, query: str, top_k: int = 5) -> List[Dict]:
        """在指定知识库中搜索最相关的文档块"""
        try:
            kb_data = self.load_kb(kb_name)
        except ValueError as e:
            print(f"加载知识库失败: {e}")
            return []
        
        if not kb_data["chunks"]:
            return []
        
        # 如果向量模型不可用，使用关键词匹配模式
        if not self.use_vector_search or self.model is None:
            return self._keyword_search(kb_data, query, top_k)
        
        try:
            # 查询嵌入
            query_embedding = self.model.encode([query])
            
            # 计算相似度
            similarities = self._calculate_similarity(query_embedding[0], kb_data["embeddings"])
            
            # 获取top_k个最相似的块
            top_indices = np.argsort(similarities)[::-1][:top_k]
            
            results = []
            for idx in top_indices:
                chunk = kb_data["chunks"][idx]
                results.append({
                    "content": chunk["content"],
                    "doc_name": chunk["doc_name"],
                    "kb_name": chunk["kb_name"],
                    "similarity": float(similarities[idx])
                })
            
            return results
        except Exception as e:
            print(f"向量搜索失败，切换到关键词匹配: {e}")
            return self._keyword_search(kb_data, query, top_k)
    
    def _keyword_search(self, kb_data: Dict, query: str, top_k: int = 5) -> List[Dict]:
        """关键词匹配搜索（向量搜索的降级方案）"""
        chunks = kb_data["chunks"]
        
        # 简单关键词匹配：检查查询词是否出现在文本中
        query_words = query.lower().split()
        
        scored_chunks = []
        for chunk in chunks:
            content_lower = chunk["content"].lower()
            score = 0
            
            # 计算匹配的关键词数量
            for word in query_words:
                if word in content_lower:
                    score += 1
                # 如果关键词较长，考虑部分匹配
                elif len(word) > 3:
                    for i in range(len(word) - 2):
                        substring = word[i:i+3]
                        if substring in content_lower:
                            score += 0.3
            
            # 考虑词频
            if score > 0:
                # 计算简单相似度（关键词匹配比例）
                similarity = min(score / len(query_words), 1.0) if query_words else 0
                scored_chunks.append((chunk, similarity))
        
        # 按分数排序
        scored_chunks.sort(key=lambda x: x[1], reverse=True)
        
        # 返回前top_k个结果
        results = []
        for chunk, similarity in scored_chunks[:top_k]:
            results.append({
                "content": chunk["content"],
                "doc_name": chunk["doc_name"],
                "kb_name": chunk["kb_name"],
                "similarity": float(similarity)
            })
        
        return results
    
    def get_all_kb_names(self) -> List[str]:
        """获取所有知识库名称"""
        if not self.kb_base_dir.exists():
            return []
        
        kb_names = []
        for item in self.kb_base_dir.iterdir():
            if item.is_dir():
                meta_file = item / "metadata.json"
                if meta_file.exists():
                    kb_names.append(item.name)
        return kb_names


def process_kb_references(prompt: str, kb_manager: KBManager) -> str:
    """
    处理知识库引用标记 {kb:知识库名称}，将其替换为检索到的内容
    """
    import re
    
    # 查找所有 {kb:知识库名称} 格式的标记
    pattern = r'\{kb:([^}]+)\}'
    
    def replace_kb_reference(match):
        kb_name = match.group(1).strip()
        
        # 提取查询词（如果有的话）
        # 格式可能是 {kb:知识库名称:查询词} 或 {kb:知识库名称}
        query_parts = kb_name.split(':', 1)
        if len(query_parts) > 1:
            kb_name = query_parts[0].strip()
            query = query_parts[1].strip()
        else:
            # 如果没有指定查询词，使用整个原始prompt作为查询
            query = prompt.replace(match.group(0), "").strip()
        
        # 在知识库中搜索相关内容
        results = kb_manager.search_kb(kb_name, query, top_k=3)
        
        if not results:
            return f"[知识库 {kb_name} 中未找到相关内容]"
        
        # 拼接最相关的几个结果
        retrieved_content = "\n".join([f"参考信息 {i+1}: {res['content']}" 
                                     for i, res in enumerate(results)])
        
        return f"\n[知识库参考信息]\n{retrieved_content}\n[结束参考信息]"
    
    # 替换所有知识库引用
    processed_prompt = re.sub(pattern, replace_kb_reference, prompt)
    
    return processed_prompt