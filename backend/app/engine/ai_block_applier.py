from __future__ import annotations

from pathlib import Path
import re
from typing import Any, Literal
from uuid import uuid4

from docx import Document
from docx.text.paragraph import Paragraph
from pydantic import BaseModel

from app.engine.ai_generator import AIGenerateResult
from app.models.enums import AIBlockStatus
from app.models.trace_models import AIBlockTrace

AI_MARKER_PATTERN = re.compile(r"§AIBLOCK\d+§")


class ApplyAIBlocksInput(BaseModel):
    rendered_docx_path: Path
    output_docx_path: Path
    ai_results: list[AIGenerateResult]
    ai_enabled: bool


class ApplyAIBlocksResult(BaseModel):
    output_docx_path: Path
    ai_status: Literal["not_used", "success", "partial_failed", "failed"]
    ai_blocks: list[AIBlockTrace]
    error_message: str = ""


def apply_ai_blocks(input_data: ApplyAIBlocksInput) -> ApplyAIBlocksResult:
    rendered_path = Path(input_data.rendered_docx_path)
    output_path = Path(input_data.output_docx_path)
    temp_output_path = output_path.with_name(f".{output_path.name}.{uuid4().hex}.tmp.docx")

    try:
        if not rendered_path.exists():
            return _failed_result(output_path, f"Rendered docx file not found: {rendered_path}")
        if not rendered_path.is_file():
            return _failed_result(output_path, f"Rendered docx path is not a file: {rendered_path}")

        output_path.parent.mkdir(parents=True, exist_ok=True)
        document = Document(str(rendered_path))
        result_by_block_id = {result.block_id: result for result in input_data.ai_results}
        ai_blocks: list[AIBlockTrace] = []

        for paragraph in _iter_all_paragraphs(document):
            markers = _find_ai_markers(paragraph.text)
            for marker in markers:
                ai_result = result_by_block_id.get(_block_id_from_marker(marker))
                ai_blocks.append(_build_ai_block_trace(marker, paragraph.text, ai_result, input_data.ai_enabled))
                replacement = _replacement_text(paragraph.text, marker, ai_result, input_data.ai_enabled)
                _replace_paragraph_text(paragraph, replacement)

        document.save(str(temp_output_path))
        Document(str(temp_output_path))
        temp_output_path.replace(output_path)

        return ApplyAIBlocksResult(
            output_docx_path=output_path,
            ai_status=_summarize_ai_status(input_data.ai_enabled, ai_blocks),
            ai_blocks=ai_blocks,
        )
    except Exception as exc:
        _remove_if_exists(temp_output_path)
        return _failed_result(output_path, f"Failed to apply AI blocks: {exc}")


def find_ai_markers(text: str) -> list[str]:
    return _find_ai_markers(text)


def _iter_all_paragraphs(document: Document) -> list[Paragraph]:
    paragraphs = list(document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                paragraphs.extend(cell.paragraphs)
    return paragraphs


def _find_ai_markers(text: str) -> list[str]:
    markers: list[str] = []
    seen: set[str] = set()
    for match in AI_MARKER_PATTERN.finditer(text):
        marker = match.group(0)
        if marker not in seen:
            markers.append(marker)
            seen.add(marker)
    return markers


def _replacement_text(
    paragraph_text: str,
    marker: str,
    ai_result: AIGenerateResult | None,
    ai_enabled: bool,
) -> str:
    original_without_marker = paragraph_text.replace(marker, "").strip()
    if not ai_enabled:
        return original_without_marker
    if ai_result is not None and ai_result.status == "success":
        generated_text = ai_result.generated_text.strip()
        if generated_text:
            return generated_text
    return original_without_marker


def _replace_paragraph_text(paragraph: Paragraph, text: str) -> None:
    first_run = paragraph.runs[0] if paragraph.runs else None
    style_snapshot = _run_style_snapshot(first_run)
    paragraph._p.clear_content()
    lines = text.splitlines() or [""]
    run = paragraph.add_run(lines[0])
    _apply_run_style_snapshot(run, style_snapshot)
    for line in lines[1:]:
        run.add_break()
        run.add_text(line)


def _run_style_snapshot(run: Any | None) -> dict[str, Any]:
    if run is None:
        return {}
    font = run.font
    return {
        "bold": run.bold,
        "italic": run.italic,
        "underline": run.underline,
        "style": run.style,
        "font_name": font.name,
        "font_size": font.size,
        "font_color": font.color.rgb if font.color is not None else None,
    }


def _apply_run_style_snapshot(run: Any, snapshot: dict[str, Any]) -> None:
    if not snapshot:
        return
    run.bold = snapshot.get("bold")
    run.italic = snapshot.get("italic")
    run.underline = snapshot.get("underline")
    if snapshot.get("style") is not None:
        run.style = snapshot["style"]
    font = run.font
    font.name = snapshot.get("font_name")
    font.size = snapshot.get("font_size")
    if snapshot.get("font_color") is not None:
        font.color.rgb = snapshot["font_color"]


def _build_ai_block_trace(
    marker: str,
    original_block_text: str,
    ai_result: AIGenerateResult | None,
    ai_enabled: bool,
) -> AIBlockTrace:
    if not ai_enabled:
        status = AIBlockStatus.NOT_USED
    elif ai_result is not None and ai_result.status == "success":
        status = AIBlockStatus.SUCCESS
    else:
        status = AIBlockStatus.FAILED

    return AIBlockTrace(
        block_id=_block_id_from_marker(marker),
        marker=marker,
        comment_id=ai_result.comment_id if ai_result is not None else None,
        selected_text=ai_result.selected_text if ai_result is not None else "",
        status=status,
        original_block_text=original_block_text,
        prompt_template=ai_result.prompt_template if ai_result is not None else "",
        prompt_rendered=ai_result.prompt_rendered if ai_result is not None else "",
        model=ai_result.model if ai_result is not None else "",
        input_variables=[],
        generated_text=ai_result.generated_text if ai_result is not None and ai_result.status == "success" else "",
        error_message=ai_result.error_message if ai_result is not None and ai_result.status == "failed" else "",
        started_at=ai_result.started_at if ai_result is not None else None,
        completed_at=ai_result.completed_at if ai_result is not None else None,
    )


def _summarize_ai_status(ai_enabled: bool, ai_blocks: list[AIBlockTrace]) -> Literal["not_used", "success", "partial_failed", "failed"]:
    if not ai_enabled:
        return "not_used"
    if not ai_blocks:
        return "success"

    success_count = sum(1 for block in ai_blocks if block.status == AIBlockStatus.SUCCESS)
    failed_count = sum(1 for block in ai_blocks if block.status == AIBlockStatus.FAILED)
    if failed_count == 0:
        return "success"
    if success_count == 0:
        return "failed"
    return "partial_failed"


def _block_id_from_marker(marker: str) -> str:
    return marker.strip("§")


def _failed_result(output_path: Path, message: str) -> ApplyAIBlocksResult:
    return ApplyAIBlocksResult(
        output_docx_path=output_path,
        ai_status="failed",
        ai_blocks=[],
        error_message=message,
    )


def _remove_if_exists(path: Path) -> None:
    try:
        path.unlink(missing_ok=True)
    except OSError:
        pass
