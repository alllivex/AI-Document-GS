from datetime import datetime, timezone
from pathlib import Path
import sys
import zipfile

from docx import Document

PROJECT_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(PROJECT_ROOT))

from app.engine.ai_block_applier import ApplyAIBlocksInput, apply_ai_blocks
from app.engine.ai_generator import AIGenerateResult
from app.engine.docx_finalizer import finalize_docx

NOW = datetime(2026, 6, 14, 12, 0, tzinfo=timezone.utc)


def make_ai_result(
    block_id: str,
    status: str,
    generated_text: str = "",
    error_message: str = "",
) -> AIGenerateResult:
    return AIGenerateResult(
        block_id=block_id,
        status=status,
        prompt_template="Analyze {{ customer_info.customer_name }}",
        prompt_rendered="Analyze Acme",
        generated_text=generated_text,
        model="deepseek-chat",
        error_message=error_message,
        started_at=NOW,
        completed_at=NOW,
    )


def write_docx(path: Path) -> None:
    document = Document()
    document.add_paragraph("客户风险分析：§AIBLOCK0§ 原始风险描述")
    document.save(path)


def write_docx_with_table_marker(path: Path) -> None:
    document = Document()
    table = document.add_table(rows=1, cols=1)
    table.cell(0, 0).text = "表格段落 §AIBLOCK0§ 原始内容"
    document.save(path)


def full_text(path: Path) -> str:
    document = Document(path)
    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.extend(paragraph.text for paragraph in cell.paragraphs)
    return "\n".join(parts)


def has_comments_part(path: Path) -> bool:
    with zipfile.ZipFile(path) as archive:
        return "word/comments.xml" in archive.namelist()


def test_ai_success_replaces_marker_paragraph_with_generated_text(tmp_path) -> None:
    rendered_path = tmp_path / "rendered.docx"
    output_path = tmp_path / "final.docx"
    write_docx(rendered_path)

    result = apply_ai_blocks(
        ApplyAIBlocksInput(
            rendered_docx_path=rendered_path,
            output_docx_path=output_path,
            ai_results=[make_ai_result("AIBLOCK0", "success", "AI生成的风险分析")],
            ai_enabled=True,
        )
    )

    assert result.error_message == ""
    assert result.ai_status == "success"
    assert output_path.exists()
    assert full_text(output_path) == "AI生成的风险分析"
    assert "§AIBLOCK" not in full_text(output_path)
    assert result.ai_blocks[0].marker == "§AIBLOCK0§"
    assert result.ai_blocks[0].status == "success"


def test_ai_failure_removes_marker_and_keeps_original_paragraph_text(tmp_path) -> None:
    rendered_path = tmp_path / "rendered.docx"
    output_path = tmp_path / "final.docx"
    write_docx(rendered_path)

    result = apply_ai_blocks(
        ApplyAIBlocksInput(
            rendered_docx_path=rendered_path,
            output_docx_path=output_path,
            ai_results=[make_ai_result("AIBLOCK0", "failed", error_message="AI generation failed")],
            ai_enabled=True,
        )
    )

    assert result.ai_status == "failed"
    assert full_text(output_path) == "客户风险分析： 原始风险描述"
    assert "§AIBLOCK" not in full_text(output_path)
    assert result.ai_blocks[0].status == "failed"
    assert result.ai_blocks[0].error_message == "AI generation failed"


def test_ai_disabled_removes_marker_without_calling_ai_result(tmp_path) -> None:
    rendered_path = tmp_path / "rendered.docx"
    output_path = tmp_path / "final.docx"
    write_docx(rendered_path)

    result = finalize_docx(
        rendered_docx_path=rendered_path,
        output_docx_path=output_path,
        ai_results=[],
        ai_enabled=False,
    )

    assert result.ai_status == "not_used"
    assert full_text(output_path) == "客户风险分析： 原始风险描述"
    assert "§AIBLOCK" not in full_text(output_path)
    assert result.ai_blocks[0].status == "not_used"


def test_ai_marker_inside_table_cell_is_removed(tmp_path) -> None:
    rendered_path = tmp_path / "rendered_table.docx"
    output_path = tmp_path / "final_table.docx"
    write_docx_with_table_marker(rendered_path)

    result = apply_ai_blocks(
        ApplyAIBlocksInput(
            rendered_docx_path=rendered_path,
            output_docx_path=output_path,
            ai_results=[make_ai_result("AIBLOCK0", "success", "表格AI文本")],
            ai_enabled=True,
        )
    )

    assert result.ai_status == "success"
    assert full_text(output_path) == "表格AI文本"
    assert "§AIBLOCK" not in full_text(output_path)


def test_final_docx_has_no_word_comments_and_can_be_opened(tmp_path) -> None:
    rendered_path = tmp_path / "rendered.docx"
    output_path = tmp_path / "final.docx"
    write_docx(rendered_path)

    result = apply_ai_blocks(
        ApplyAIBlocksInput(
            rendered_docx_path=rendered_path,
            output_docx_path=output_path,
            ai_results=[make_ai_result("AIBLOCK0", "success", "AI生成的风险分析")],
            ai_enabled=True,
        )
    )

    assert result.ai_status == "success"
    opened = Document(output_path)
    assert opened.paragraphs[0].text == "AI生成的风险分析"
    assert not has_comments_part(output_path)
