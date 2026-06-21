from io import BytesIO
from pathlib import Path
import sqlite3
import sys

from docx import Document
from fastapi.testclient import TestClient
from openpyxl import Workbook

PROJECT_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(PROJECT_ROOT))

from app.core.config import get_settings
from app.db.connection import get_connection
from app.db.init_db import init_db


def reset_settings(tmp_path, monkeypatch):
    monkeypatch.setenv("PROJECT_ROOT", str(tmp_path))
    monkeypatch.setenv("DATABASE_URL", f"sqlite:///{(tmp_path / 'config' / 'db.sqlite').as_posix()}")
    get_settings.cache_clear()
    return get_settings()


def docx_bytes(text: str = "Hello {{ customer_info.customer_name }}") -> bytes:
    buffer = BytesIO()
    document = Document()
    document.add_paragraph(text)
    document.save(buffer)
    return buffer.getvalue()


def docx_text(file_bytes: bytes) -> str:
    document = Document(BytesIO(file_bytes))
    return "\n".join(paragraph.text for paragraph in document.paragraphs)


def xlsx_bytes(text: str = "{{ customer_info.customer_name }}") -> bytes:
    buffer = BytesIO()
    workbook = Workbook()
    workbook.active["A1"] = text
    workbook.save(buffer)
    return buffer.getvalue()


def test_template_file_upload_and_download_xlsx(tmp_path, monkeypatch) -> None:
    settings = reset_settings(tmp_path, monkeypatch)
    init_db(settings.database_path)
    from app.main import create_app

    with TestClient(create_app()) as client:
        response = client.post(
            "/api/settings/template-files",
            data={"template_name": "Excel Report", "description": ""},
            files={"file": ("report.xlsx", xlsx_bytes(), "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")},
        )
        assert response.status_code == 200
        payload = response.json()["data"]
        assert payload["template_file_type"] == "xlsx"
        assert payload["template_file"].endswith(".xlsx")
        assert (settings.templates_dir / payload["template_file"]).is_file()

        download = client.get(f"/api/settings/template-files/{payload['template_id']}/download")
        assert download.status_code == 200
        assert download.headers["content-type"] == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"


def assert_success(payload: dict) -> None:
    assert payload["success"] is True
    assert payload["request_id"]
    assert payload["timestamp"]


def seed_template_requirements(connection: sqlite3.Connection, template_id: int) -> None:
    now = "2026-06-14T12:00:00+00:00"
    connection.execute(
        """
        UPDATE templates
        SET main_table = ?
        WHERE id = ?
        """,
        ("customer_info", template_id),
    )
    connection.execute(
        """
        INSERT INTO template_tables (
            template_id, table_name, table_name_cn, role, relation_type,
            main_join_key, table_join_key, required, sort_order, description,
            created_at, updated_at
        )
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (template_id, "customer_info", "Customer Info", "main", "main", "", "", 1, 0, "", now, now),
    )
    connection.executemany(
        """
        INSERT INTO fields (
            table_name, table_name_cn, field_name, field_name_cn, data_type,
            is_primary_key, required, display_format, description, created_at, updated_at
        )
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        [
            ("customer_info", "Customer Info", "customer_id", "Customer ID", "string", 1, 1, "", "", now, now),
            ("customer_info", "Customer Info", "customer_name", "Customer Name", "string", 0, 1, "", "", now, now),
        ],
    )


def upload_template(client: TestClient, name: str = "Branch Report", content: bytes | None = None) -> int:
    response = client.post(
        "/api/settings/template-files",
        data={"template_name": name, "description": "Template description"},
        files={
            "file": (
                f"{name}.docx",
                content or docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    assert response.status_code == 200
    payload = response.json()
    assert_success(payload)
    return int(payload["data"]["template_id"])


def test_template_file_upload_list_download_deactivate_and_activate(tmp_path, monkeypatch) -> None:
    settings = reset_settings(tmp_path, monkeypatch)
    init_db(settings.database_path)

    from app.main import create_app

    with TestClient(create_app()) as client:
        template_id = upload_template(client)
        stored_path = settings.templates_dir / f"template_{template_id}.docx"
        assert stored_path.is_file()

        list_response = client.get("/api/settings/template-files")
        assert list_response.status_code == 200
        list_payload = list_response.json()
        assert_success(list_payload)
        assert list_payload["data"]["total"] == 1
        assert list_payload["data"]["items"][0]["template_name"] == "Branch Report"

        center_response = client.get("/api/templates")
        assert center_response.status_code == 200
        assert center_response.json()["data"]["total"] == 1

        download_response = client.get(f"/api/settings/template-files/{template_id}/download")
        assert download_response.status_code == 200
        assert download_response.headers["content-type"] == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        assert download_response.content.startswith(b"PK")

        deactivate_response = client.delete(f"/api/settings/template-files/{template_id}")
        assert deactivate_response.status_code == 200
        assert deactivate_response.json()["data"]["is_active"] is False
        assert stored_path.is_file()

        active_center_response = client.get("/api/templates")
        assert active_center_response.status_code == 200
        assert active_center_response.json()["data"]["total"] == 0

        activate_response = client.post(f"/api/settings/template-files/{template_id}/activate")
        assert activate_response.status_code == 200
        assert activate_response.json()["data"]["is_active"] is True

        reactivated_center_response = client.get("/api/templates")
        assert reactivated_center_response.status_code == 200
        assert reactivated_center_response.json()["data"]["total"] == 1

    with sqlite3.connect(settings.database_path) as connection:
        row = connection.execute(
            """
            SELECT template_name, template_file, original_filename, description, is_active
            FROM templates
            WHERE id = ?
            """,
            (template_id,),
        ).fetchone()
    assert row == ("Branch Report", f"template_{template_id}.docx", "Branch Report.docx", "Template description", 1)


def test_template_file_replace_valid_docx_updates_existing_file(tmp_path, monkeypatch) -> None:
    settings = reset_settings(tmp_path, monkeypatch)
    init_db(settings.database_path)

    from app.main import create_app

    with TestClient(create_app()) as client:
        template_id = upload_template(client, content=docx_bytes("Old {{ customer_info.customer_name }}"))
        with get_connection(settings.database_path) as connection:
            seed_template_requirements(connection, template_id)

        response = client.put(
            f"/api/settings/template-files/{template_id}/file",
            files={
                "file": (
                    "replacement.docx",
                    docx_bytes("New {{ customer_info.customer_name }}"),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

        assert response.status_code == 200
        payload = response.json()
        assert_success(payload)
        assert payload["data"]["template_id"] == template_id
        assert payload["data"]["template_file"] == f"template_{template_id}.docx"
        assert payload["data"]["original_filename"] == "replacement.docx"

        download_response = client.get(f"/api/settings/template-files/{template_id}/download")
        assert download_response.status_code == 200
        assert "New {{ customer_info.customer_name }}" in docx_text(download_response.content)


def test_template_file_replace_rejects_non_docx(tmp_path, monkeypatch) -> None:
    settings = reset_settings(tmp_path, monkeypatch)
    init_db(settings.database_path)

    from app.main import create_app

    with TestClient(create_app()) as client:
        template_id = upload_template(client)
        response = client.put(
            f"/api/settings/template-files/{template_id}/file",
            files={"file": ("bad.txt", b"not docx", "text/plain")},
        )

    payload = response.json()
    assert response.status_code == 400
    assert payload["success"] is False
    assert payload["error"]["code"] == "UPLOAD_FILE_INVALID"


def test_template_file_replace_rejects_invalid_variables_and_keeps_old_file(tmp_path, monkeypatch) -> None:
    settings = reset_settings(tmp_path, monkeypatch)
    init_db(settings.database_path)

    from app.main import create_app

    with TestClient(create_app()) as client:
        old_content = docx_bytes("Old {{ customer_info.customer_name }}")
        template_id = upload_template(client, content=old_content)
        with get_connection(settings.database_path) as connection:
            seed_template_requirements(connection, template_id)

        response = client.put(
            f"/api/settings/template-files/{template_id}/file",
            files={
                "file": (
                    "bad_variables.docx",
                    docx_bytes("Bad {{ data.customer_name }}"),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

        payload = response.json()
        assert response.status_code == 400
        assert payload["success"] is False
        assert payload["error"]["code"] == "VALIDATION_FAILED"

        download_response = client.get(f"/api/settings/template-files/{template_id}/download")
        assert download_response.status_code == 200
        assert "Old {{ customer_info.customer_name }}" in docx_text(download_response.content)


def test_template_file_upload_rejects_non_docx(tmp_path, monkeypatch) -> None:
    settings = reset_settings(tmp_path, monkeypatch)
    init_db(settings.database_path)

    from app.main import create_app

    with TestClient(create_app()) as client:
        response = client.post(
            "/api/settings/template-files",
            data={"template_name": "Bad Template", "description": ""},
            files={"file": ("bad.txt", b"not docx", "text/plain")},
        )

    payload = response.json()
    assert response.status_code == 400
    assert payload["success"] is False
    assert payload["error"]["code"] == "UPLOAD_FILE_INVALID"


def test_template_file_upload_rejects_invalid_docx_content(tmp_path, monkeypatch) -> None:
    settings = reset_settings(tmp_path, monkeypatch)
    init_db(settings.database_path)

    from app.main import create_app

    with TestClient(create_app()) as client:
        response = client.post(
            "/api/settings/template-files",
            data={"template_name": "Broken Template", "description": ""},
            files={
                "file": (
                    "broken.docx",
                    b"not a valid docx",
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )

    payload = response.json()
    assert response.status_code == 400
    assert payload["success"] is False
    assert payload["error"]["code"] == "UPLOAD_FILE_INVALID"


def test_template_file_api_adds_metadata_columns_for_existing_schema(tmp_path, monkeypatch) -> None:
    settings = reset_settings(tmp_path, monkeypatch)
    init_db(settings.database_path)

    from app.main import create_app

    with TestClient(create_app()) as client:
        response = client.get("/api/settings/template-files")

    assert response.status_code == 200
    assert response.json()["data"] == {"items": [], "total": 0}
    with get_connection(settings.database_path) as connection:
        columns = {row["name"] for row in connection.execute("PRAGMA table_info(templates)").fetchall()}
    assert "original_filename" in columns
    assert "description" in columns
