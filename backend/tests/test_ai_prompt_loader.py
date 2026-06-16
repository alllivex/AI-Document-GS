from pathlib import Path
import sys
import zipfile

from docx import Document
from lxml import etree

PROJECT_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(PROJECT_ROOT))

from app.engine.ai_prompt_loader import (
    AIPromptLoadInput,
    analyze_ai_prompt_bindings,
    extract_prompt_from_comment_text,
    get_ai_block_markers_for_template,
    load_ai_prompts,
    load_ai_prompts_from_template,
)

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
CT_NS = "http://schemas.openxmlformats.org/package/2006/content-types"
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
W = f"{{{W_NS}}}"


def write_template_with_ai_comments(path: Path, comments: list[tuple[str, str, str]]) -> None:
    document = Document()
    document.add_paragraph("Intro §AIBLOCK0§")
    document.add_paragraph("Next §AIBLOCK1§")
    document.add_paragraph("Final §AIBLOCK2§")
    document.save(path)
    add_comments_xml(path, comments)


def add_comments_xml(path: Path, comments: list[tuple[str, str, str]]) -> None:
    with zipfile.ZipFile(path, "r") as archive:
        files = {name: archive.read(name) for name in archive.namelist()}

    files["word/document.xml"] = _add_comment_ranges(files["word/document.xml"], comments)
    files["word/comments.xml"] = _build_comments_xml(comments)
    files["[Content_Types].xml"] = _add_comments_content_type(files["[Content_Types].xml"])
    files["word/_rels/document.xml.rels"] = _add_comments_relationship(files["word/_rels/document.xml.rels"])

    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as archive:
        for name, data in files.items():
            archive.writestr(name, data)


def _add_comment_ranges(document_xml: bytes, comments: list[tuple[str, str, str]]) -> bytes:
    root = etree.fromstring(document_xml)
    for comment_id, _comment_text, marker in comments:
        paragraph = _find_paragraph_containing(root, marker)
        runs = paragraph.findall(f"{W}r")
        start = etree.Element(f"{W}commentRangeStart")
        start.set(f"{W}id", comment_id)
        end = etree.Element(f"{W}commentRangeEnd")
        end.set(f"{W}id", comment_id)
        reference_run = etree.Element(f"{W}r")
        reference = etree.SubElement(reference_run, f"{W}commentReference")
        reference.set(f"{W}id", comment_id)

        paragraph.insert(paragraph.index(runs[0]), start)
        paragraph.insert(paragraph.index(runs[-1]) + 1, end)
        paragraph.insert(paragraph.index(end) + 1, reference_run)

    return etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)


def _find_paragraph_containing(root, marker: str):
    for paragraph in root.xpath(".//w:p", namespaces={"w": W_NS}):
        text = "".join(paragraph.xpath(".//w:t/text()", namespaces={"w": W_NS}))
        if marker in text:
            return paragraph
    raise AssertionError(f"marker not found in test document: {marker}")


def _build_comments_xml(comments: list[tuple[str, str, str]]) -> bytes:
    comments_root = etree.Element(f"{W}comments", nsmap={"w": W_NS})
    for comment_id, comment_text, _marker in comments:
        comment = etree.SubElement(comments_root, f"{W}comment")
        comment.set(f"{W}id", comment_id)
        paragraph = etree.SubElement(comment, f"{W}p")
        run = etree.SubElement(paragraph, f"{W}r")
        text = etree.SubElement(run, f"{W}t")
        text.text = comment_text
    return etree.tostring(comments_root, xml_declaration=True, encoding="UTF-8", standalone=True)


def _add_comments_content_type(content_types_xml: bytes) -> bytes:
    content_types = etree.fromstring(content_types_xml)
    override_tag = f"{{{CT_NS}}}Override"
    if not any(element.get("PartName") == "/word/comments.xml" for element in content_types.findall(override_tag)):
        override = etree.Element(override_tag)
        override.set("PartName", "/word/comments.xml")
        override.set("ContentType", "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml")
        content_types.append(override)
    return etree.tostring(content_types, xml_declaration=True, encoding="UTF-8", standalone=True)


def _add_comments_relationship(rels_xml: bytes) -> bytes:
    relationships = etree.fromstring(rels_xml)
    relationship_tag = f"{{{REL_NS}}}Relationship"
    if not any(element.get("Target") == "comments.xml" for element in relationships.findall(relationship_tag)):
        relationship = etree.Element(relationship_tag)
        relationship.set("Id", "rId999")
        relationship.set("Type", "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments")
        relationship.set("Target", "comments.xml")
        relationships.append(relationship)
    return etree.tostring(relationships, xml_declaration=True, encoding="UTF-8", standalone=True)


def test_get_ai_block_markers_for_template_reads_markers_in_document_order(tmp_path) -> None:
    template_path = tmp_path / "template.docx"
    write_template_with_ai_comments(
        template_path,
        [
            ("0", 'prompt="Analyze {{ customer_info.customer_name }}"', "§AIBLOCK0§"),
            ("1", 'prompt="Summarize loan {{ loan_summary.loan_balance }}"', "§AIBLOCK1§"),
        ],
    )

    markers = get_ai_block_markers_for_template(template_path)

    assert markers == ["§AIBLOCK0§", "§AIBLOCK1§", "§AIBLOCK2§"]


def test_load_ai_prompts_from_template_binds_prompts_by_word_comment_anchor(tmp_path) -> None:
    template_path = tmp_path / "template.docx"
    write_template_with_ai_comments(
        template_path,
        [
            ("9", 'prompt="Summarize loan {{ loan_summary.loan_balance }}"', "§AIBLOCK1§"),
            ("3", 'prompt="Analyze {{ customer_info.customer_name }}"', "§AIBLOCK0§"),
            ("7", 'prompt="Write final note"', "§AIBLOCK2§"),
        ],
    )

    prompts = load_ai_prompts_from_template(template_path)

    assert len(prompts) == 3
    prompt_by_block_id = {prompt.block_id: prompt for prompt in prompts}
    assert prompt_by_block_id["AIBLOCK0"].comment_id == "3"
    assert prompt_by_block_id["AIBLOCK0"].prompt_template == "Analyze {{ customer_info.customer_name }}"
    assert prompt_by_block_id["AIBLOCK0"].selected_text == "Intro §AIBLOCK0§"
    assert prompt_by_block_id["AIBLOCK1"].comment_id == "9"
    assert prompt_by_block_id["AIBLOCK1"].prompt_template == "Summarize loan {{ loan_summary.loan_balance }}"
    assert prompt_by_block_id["AIBLOCK2"].comment_id == "7"


def test_load_ai_prompts_accepts_input_model(tmp_path) -> None:
    template_path = tmp_path / "template.docx"
    write_template_with_ai_comments(
        template_path,
        [
            ("0", 'prompt="Analyze {{ customer_info.customer_name }}"', "§AIBLOCK0§"),
            ("1", 'prompt="Summarize loan {{ loan_summary.loan_balance }}"', "§AIBLOCK1§"),
            ("2", 'prompt="Write final note"', "§AIBLOCK2§"),
        ],
    )

    prompts = load_ai_prompts(AIPromptLoadInput(template_path=template_path))

    assert [prompt.marker for prompt in prompts] == ["§AIBLOCK0§", "§AIBLOCK1§", "§AIBLOCK2§"]


def test_load_ai_prompts_reports_marker_without_prompt_comment(tmp_path) -> None:
    template_path = tmp_path / "template.docx"
    write_template_with_ai_comments(
        template_path,
        [
            ("0", 'prompt="Analyze {{ customer_info.customer_name }}"', "§AIBLOCK0§"),
        ],
    )

    analysis = analyze_ai_prompt_bindings(template_path)

    assert {issue.code for issue in analysis.issues} == {"ai_marker_missing_prompt_comment"}
    assert {issue.marker for issue in analysis.issues} == {"§AIBLOCK1§", "§AIBLOCK2§"}


def test_load_ai_prompts_reports_prompt_comment_without_marker(tmp_path) -> None:
    template_path = tmp_path / "template.docx"
    write_template_with_ai_comments(
        template_path,
        [
            ("0", 'prompt="Analyze {{ customer_info.customer_name }}"', "§AIBLOCK0§"),
            ("1", 'prompt="No marker here"', "§AIBLOCK1§"),
            ("2", 'prompt="Write final note"', "§AIBLOCK2§"),
        ],
    )

    with zipfile.ZipFile(template_path, "r") as archive:
        files = {name: archive.read(name) for name in archive.namelist()}
    root = etree.fromstring(files["word/document.xml"])
    for element in root.xpath(".//w:commentRangeStart[@w:id='1'] | .//w:commentRangeEnd[@w:id='1']", namespaces={"w": W_NS}):
        element.getparent().remove(element)
    files["word/document.xml"] = etree.tostring(root, xml_declaration=True, encoding="UTF-8", standalone=True)
    with zipfile.ZipFile(template_path, "w", zipfile.ZIP_DEFLATED) as archive:
        for name, data in files.items():
            archive.writestr(name, data)

    analysis = analyze_ai_prompt_bindings(template_path)

    assert "prompt_comment_without_ai_marker" in {issue.code for issue in analysis.issues}


def test_load_ai_prompts_reports_comment_covering_multiple_markers(tmp_path) -> None:
    template_path = tmp_path / "template.docx"
    document = Document()
    document.add_paragraph("Both §AIBLOCK0§ and §AIBLOCK1§")
    document.save(template_path)
    add_comments_xml(template_path, [("0", 'prompt="Analyze both"', "§AIBLOCK0§")])

    analysis = analyze_ai_prompt_bindings(template_path)

    assert analysis.issues[0].code == "prompt_comment_covers_multiple_ai_markers"


def test_load_ai_prompts_reports_duplicate_marker(tmp_path) -> None:
    template_path = tmp_path / "template.docx"
    document = Document()
    document.add_paragraph("First §AIBLOCK0§")
    document.add_paragraph("Second §AIBLOCK0§")
    document.save(template_path)
    add_comments_xml(template_path, [("0", 'prompt="Analyze"', "§AIBLOCK0§")])

    analysis = analyze_ai_prompt_bindings(template_path)

    assert "duplicate_ai_marker" in {issue.code for issue in analysis.issues}


def test_extract_prompt_from_comment_text_requires_prompt_attribute() -> None:
    assert extract_prompt_from_comment_text('prefix prompt="Hello {{ name }}" suffix') == "Hello {{ name }}"
