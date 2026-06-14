from __future__ import annotations

from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path
import zipfile
import sys

from docx import Document
from fastapi.testclient import TestClient
import pandas as pd

PROJECT_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(PROJECT_ROOT))

from app.core.config import get_settings
from app.db.init_db import init_db
from app.db.connection import get_connection

NOW = datetime(2026, 6, 14, 12, 0, tzinfo=timezone.utc).isoformat()


def reset_settings(tmp_path: Path, monkeypatch):
    monkeypatch.setenv("PROJECT_ROOT", str(tmp_path))
    monkeypatch.setenv("DATABASE_URL", f"sqlite:///{(tmp_path / 'config' / 'db.sqlite').as_posix()}")
    monkeypatch.delenv("DEEPSEEK_API_KEY", raising=False)
    get_settings.cache_clear()
    return get_settings()


def seed_template(connection) -> None:
    connection.execute(
        """
        INSERT INTO templates (id, template_name, template_file, template_path, main_table, created_at, updated_at)
        VALUES (1, 'due diligence', 'due_diligence.docx', 'templates/due_diligence.docx', 'customer_info', ?, ?)
        """,
        (NOW, NOW),
    )
    connection.executemany(
        """
        INSERT INTO template_tables (
            template_id, table_name, table_name_cn, role, relation_type,
            main_join_key, table_join_key, required, sort_order, created_at, updated_at
        )
        VALUES (1, ?, '', ?, ?, ?, ?, 1, ?, ?, ?)
        """,
        [
            ("customer_info", "main", "main", "", "", 0, NOW, NOW),
            ("loan_summary", "aux", "one_to_one", "customer_id", "customer_id", 1, NOW, NOW),
        ],
    )
    connection.executemany(
        """
        INSERT INTO fields (
            table_name, field_name, data_type, is_primary_key, required, display_format, created_at, updated_at
        )
        VALUES (?, ?, ?, ?, 1, ?, ?, ?)
        """,
        [
            ("customer_info", "customer_id", "string", 1, "", NOW, NOW),
            ("customer_info", "customer_name", "string", 0, "", NOW, NOW),
            ("loan_summary", "customer_id", "string", 0, "", NOW, NOW),
            ("loan_summary", "loan_balance", "amount", 0, "amount_2", NOW, NOW),
        ],
    )


def write_template(templates_dir: Path) -> None:
    templates_dir.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.add_paragraph("Customer: {{ customer_info.customer_name }}")
    document.add_paragraph("Balance: {{ loan_summary.loan_balance }}")
    document.save(templates_dir / "due_diligence.docx")


def excel_bytes(rows: list[dict]) -> bytes:
    buffer = BytesIO()
    pd.DataFrame(rows).to_excel(buffer, index=False)
    return buffer.getvalue()


def assert_success(payload: dict) -> None:
    assert payload["success"] is True
    assert "request_id" in payload
    assert "timestamp" in payload
    assert payload["error"] is None


def test_api_e2e_generates_outputs_and_serves_preview_trace_downloads(tmp_path, monkeypatch) -> None:
    settings = reset_settings(tmp_path, monkeypatch)
    init_db(settings.database_path)
    write_template(settings.templates_dir)
    with get_connection(settings.database_path) as connection:
        seed_template(connection)

    from app.main import create_app

    with TestClient(create_app()) as client:
        templates_response = client.get("/api/templates")
        assert templates_response.status_code == 200
        assert_success(templates_response.json())
        assert templates_response.json()["data"]["total"] == 1

        template_response = client.get("/api/templates/1")
        assert template_response.status_code == 200
        assert_success(template_response.json())

        requirements_response = client.get("/api/templates/1/requirements")
        assert requirements_response.status_code == 200
        assert_success(requirements_response.json())

        create_response = client.post(
            "/api/tasks",
            json={"task_name": "task", "template_id": 1, "ai_enabled": False},
        )
        assert create_response.status_code == 200
        assert_success(create_response.json())
        task_id = create_response.json()["data"]["task_id"]

        customer_file = excel_bytes(
            [
                {"customer_id": "C001", "customer_name": "Acme"},
                {"customer_id": "C002", "customer_name": "Beta"},
                {"customer_id": "C003", "customer_name": "Gamma"},
            ]
        )
        loan_file = excel_bytes(
            [
                {"customer_id": "C001", "loan_balance": 1200},
                {"customer_id": "C002", "loan_balance": 800},
                {"customer_id": "C003", "loan_balance": 1500},
            ]
        )
        for table_name, file_bytes in (("customer_info", customer_file), ("loan_summary", loan_file)):
            upload_response = client.post(
                f"/api/tasks/{task_id}/upload",
                data={"table_name": table_name},
                files={"file": (f"{table_name}.xlsx", file_bytes, "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")},
            )
            assert upload_response.status_code == 200
            assert_success(upload_response.json())

        validate_response = client.post(f"/api/tasks/{task_id}/validate")
        assert validate_response.status_code == 200
        assert_success(validate_response.json())
        assert validate_response.json()["data"]["status"] == "passed"

        generate_response = client.post(f"/api/tasks/{task_id}/generate")
        assert generate_response.status_code == 200
        assert_success(generate_response.json())
        assert generate_response.json()["data"]["success_count"] == 3

        tasks_response = client.get("/api/tasks")
        assert tasks_response.status_code == 200
        assert_success(tasks_response.json())

        outputs_response = client.get(f"/api/tasks/{task_id}/outputs")
        assert outputs_response.status_code == 200
        assert_success(outputs_response.json())
        outputs = outputs_response.json()["data"]["items"]
        assert len(outputs) == 3
        doc_id = outputs[0]["doc_id"]

        preview_response = client.get(f"/api/documents/{doc_id}/preview")
        assert preview_response.status_code == 200
        assert_success(preview_response.json())
        trace_id = next(
            run["trace_id"]
            for block in preview_response.json()["data"]["blocks"]
            if block["type"] == "paragraph"
            for run in block["runs"]
            if run["trace_id"]
        )

        trace_response = client.get(f"/api/trace/{trace_id}")
        assert trace_response.status_code == 200
        assert_success(trace_response.json())
        assert trace_response.json()["data"]["trace_id"] == trace_id

        download_response = client.get(f"/api/documents/{doc_id}/download")
        assert download_response.status_code == 200
        assert download_response.headers["content-type"] == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        assert "filename=" in download_response.headers["content-disposition"]

        zip_response = client.get(f"/api/tasks/{task_id}/download-zip")
        assert zip_response.status_code == 200
        assert zip_response.headers["content-type"] == "application/zip"
        with zipfile.ZipFile(BytesIO(zip_response.content)) as archive:
            assert len(archive.namelist()) == 3


def test_api_errors_use_unified_json_structure(tmp_path, monkeypatch) -> None:
    settings = reset_settings(tmp_path, monkeypatch)
    init_db(settings.database_path)

    from app.main import create_app

    with TestClient(create_app()) as client:
        task_response = client.get("/api/tasks/missing/outputs")
        assert task_response.status_code == 404
        task_payload = task_response.json()
        assert task_payload["success"] is False
        assert task_payload["error"]["code"] == "TASK_NOT_FOUND"
        assert "request_id" in task_payload
        assert "timestamp" in task_payload

        document_response = client.get("/api/documents/missing/preview")
        assert document_response.status_code == 404
        document_payload = document_response.json()
        assert document_payload["success"] is False
        assert document_payload["error"]["code"] == "DOCUMENT_NOT_FOUND"
