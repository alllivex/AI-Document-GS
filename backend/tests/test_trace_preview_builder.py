from datetime import datetime, timezone
from pathlib import Path
import json
import sys
import zipfile

from docx import Document

PROJECT_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(PROJECT_ROOT))

from app.engine.trace_builder import BuildTracePreviewInput, build_trace_and_preview
from app.models.enums import DataType, RelationType
from app.models.trace_models import TraceItem

NOW = datetime(2026, 6, 14, 12, 0, tzinfo=timezone.utc)


def make_trace_item(
    trace_id: str,
    var_path: str,
    table_name: str,
    field_name: str,
    display_value: str,
    occurrence_index: int = 0,
) -> TraceItem:
    return TraceItem(
        trace_id=trace_id,
        var_path=var_path,
        table_name=table_name,
        table_name_cn="",
        field_name=field_name,
        field_name_cn="",
        source_file=f"{table_name}.xlsx",
        source_file_path=f"tasks/task_001/data/{table_name}.xlsx",
        pk_field="customer_id",
        pk_value="C001",
        row_index=occurrence_index,
        excel_row_number=occurrence_index + 2,
        column_index=1,
        excel_column_letter="B",
        raw_value=display_value,
        display_value=display_value,
        value_type=DataType.STRING,
        display_format="",
        occurrence_index=occurrence_index,
        source_relation_type=RelationType.MAIN,
        created_at=NOW,
    )


def write_preview_docx(path: Path) -> None:
    document = Document()
    document.add_paragraph("Customer: Acme")
    document.add_paragraph("This unmatched paragraph is still visible.")
    table = document.add_table(rows=3, cols=2)
    table.rows[0].cells[0].text = "Collateral"
    table.rows[0].cells[1].text = "Amount"
    table.rows[1].cells[0].text = "Factory"
    table.rows[1].cells[1].text = "3000.00"
    table.rows[2].cells[0].text = "Unmatched cell"
    table.rows[2].cells[1].text = "No trace"
    document.save(path)


def write_condition_docx(path: Path) -> None:
    document = Document()
    document.add_paragraph("清偿能力判断：无法清偿。")
    document.save(path)


def write_condition_template(path: Path) -> None:
    with zipfile.ZipFile(path, "w") as archive:
        archive.writestr(
            "word/document.xml",
            "<document><body><p>{% if loan_summary.house_eval_amount &lt; loan_summary.loan_balance %}"
            "无法清偿{% else %}可清偿{% endif %}</p></body></document>",
        )


def write_loop_docx(path: Path) -> None:
    document = Document()
    table = document.add_table(rows=3, cols=2)
    table.rows[0].cells[0].text = "产品名称"
    table.rows[0].cells[1].text = "不良率(%)"
    table.rows[1].cells[0].text = "产品A"
    table.rows[1].cells[1].text = "1.20%"
    table.rows[2].cells[0].text = "产品B"
    table.rows[2].cells[1].text = "2.40%"
    document.save(path)


def write_loop_template(path: Path) -> None:
    with zipfile.ZipFile(path, "w") as archive:
        archive.writestr(
            "word/document.xml",
            "<document><body><tbl>{%tr for p in products %}{{ p.product_name }} {{ p.overdue_rate }}"
            "{%tr endfor %}</tbl></body></document>",
        )


def load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8"))


def collect_preview_trace_ids(preview_payload: dict) -> list[str]:
    trace_ids: list[str] = []
    for block in preview_payload["blocks"]:
        if block["type"] == "paragraph":
            trace_ids.extend(run["trace_id"] for run in block["runs"] if run.get("trace_id"))
        elif block["type"] == "table":
            trace_ids.extend(cell["trace_id"] for cell in block["headers"] if cell.get("trace_id"))
            for row in block["rows"]:
                trace_ids.extend(cell["trace_id"] for cell in row if cell.get("trace_id"))
    return trace_ids


def build_input(tmp_path: Path, trace_map: dict[str, list[TraceItem]]) -> BuildTracePreviewInput:
    output_path = tmp_path / "due_diligence_C001.docx"
    write_preview_docx(output_path)
    return BuildTracePreviewInput(
        task_id="task_001",
        doc_id="doc_001",
        template_id=1,
        template_name="due diligence",
        template_file="due_diligence.docx",
        output_file=output_path.name,
        output_path=output_path,
        main_table="customer_info",
        primary_key_field="customer_id",
        primary_key_value="C001",
        trace_map=trace_map,
        final_docx_path=output_path,
    )


def test_build_trace_and_preview_writes_sidecar_json_files(tmp_path) -> None:
    trace_map = {
        "customer_info.customer_name": [
            make_trace_item("trace_customer_name", "customer_info.customer_name", "customer_info", "customer_name", "Acme")
        ],
        "collateral_info.collateral_name": [
            make_trace_item("trace_collateral_name", "collateral_info.collateral_name", "collateral_info", "collateral_name", "Factory")
        ],
    }

    result = build_trace_and_preview(build_input(tmp_path, trace_map))

    assert result.trace_file_path == tmp_path / "due_diligence_C001.trace.json"
    assert result.preview_file_path == tmp_path / "due_diligence_C001.preview.json"
    assert result.trace_count == 2
    assert result.ai_block_count == 0
    assert result.trace_file_path.exists()
    assert result.preview_file_path.exists()


def test_preview_clickable_text_references_trace_json_items(tmp_path) -> None:
    trace_map = {
        "customer_info.customer_name": [
            make_trace_item("trace_customer_name", "customer_info.customer_name", "customer_info", "customer_name", "Acme")
        ],
        "collateral_info.collateral_name": [
            make_trace_item("trace_collateral_name", "collateral_info.collateral_name", "collateral_info", "collateral_name", "Factory")
        ],
    }

    result = build_trace_and_preview(build_input(tmp_path, trace_map))
    trace_payload = load_json(result.trace_file_path)
    preview_payload = load_json(result.preview_file_path)

    trace_ids = {item["trace_id"] for item in trace_payload["trace_items"]}
    preview_trace_ids = set(collect_preview_trace_ids(preview_payload))

    assert preview_trace_ids == {"trace_customer_name", "trace_collateral_name"}
    assert preview_trace_ids <= trace_ids
    customer_trace = next(item for item in trace_payload["trace_items"] if item["trace_id"] == "trace_customer_name")
    assert customer_trace["source_file"] == "customer_info.xlsx"
    assert customer_trace["table_name"] == "customer_info"
    assert customer_trace["field_name"] == "customer_name"
    assert customer_trace["original_var_path"] == "customer_info.customer_name"
    assert customer_trace["canonical_var_path"] == "customer_info.customer_name"
    assert customer_trace["excel_row_number"] == 2
    assert customer_trace["raw_value"] == "Acme"
    assert customer_trace["display_value"] == "Acme"


def test_unmatched_preview_text_is_kept_without_trace_id(tmp_path) -> None:
    trace_map = {
        "customer_info.customer_name": [
            make_trace_item("trace_customer_name", "customer_info.customer_name", "customer_info", "customer_name", "Acme")
        ]
    }

    result = build_trace_and_preview(build_input(tmp_path, trace_map))
    preview_payload = load_json(result.preview_file_path)

    unmatched_block = next(
        block
        for block in preview_payload["blocks"]
        if block["type"] == "paragraph" and block["runs"][0]["text"] == "This unmatched paragraph is still visible."
    )
    assert unmatched_block["runs"] == [
        {
            "text": "This unmatched paragraph is still visible.",
            "trace_id": None,
            "trace_kind": None,
            "ai_block_id": None,
            "style": None,
        }
    ]

    table_block = next(block for block in preview_payload["blocks"] if block["type"] == "table")
    assert table_block["rows"][1][0] == {"text": "Unmatched cell", "trace_id": None, "trace_kind": None, "ai_block_id": None}


def test_repeated_display_values_use_trace_map_order(tmp_path) -> None:
    output_path = tmp_path / "repeat.docx"
    document = Document()
    document.add_paragraph("Acme / Acme / Acme")
    document.save(output_path)
    first = make_trace_item("trace_first", "customer_info.customer_name", "customer_info", "customer_name", "Acme", 0)
    second = make_trace_item("trace_second", "related.customer_name", "related", "customer_name", "Acme", 1)

    result = build_trace_and_preview(
        BuildTracePreviewInput(
            task_id="task_001",
            doc_id="doc_001",
            template_id=1,
            template_name="due diligence",
            template_file="due_diligence.docx",
            output_file=output_path.name,
            output_path=output_path,
            main_table="customer_info",
            primary_key_field="customer_id",
            primary_key_value="C001",
            trace_map={
                "customer_info.customer_name": [first],
                "related.customer_name": [second],
            },
            final_docx_path=output_path,
        )
    )

    preview_payload = load_json(result.preview_file_path)
    runs = preview_payload["blocks"][0]["runs"]
    assert [(run["text"], run["trace_id"]) for run in runs] == [
        ("Acme", "trace_first"),
        (" / ", None),
        ("Acme", "trace_second"),
        (" / ", None),
        ("Acme", "trace_second"),
    ]


def test_single_source_value_stays_clickable_when_repeated_in_paragraphs_and_table(tmp_path) -> None:
    output_path = tmp_path / "repeat_single_source.docx"
    document = Document()
    document.add_paragraph("Name: Alice, ID: ID-001, Address: ADDR-001")
    document.add_paragraph("Repeated ID: ID-001")
    document.add_paragraph("Repeated Address: ADDR-001")
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "ID"
    table.rows[0].cells[1].text = "Address"
    table.rows[1].cells[0].text = "ID-001"
    table.rows[1].cells[1].text = "ADDR-001"
    document.save(output_path)

    id_item = make_trace_item("trace_id_no", "customer.id_no", "customer", "id_no", "ID-001")
    address_item = make_trace_item("trace_address", "customer.address", "customer", "address", "ADDR-001")

    result = build_trace_and_preview(
        BuildTracePreviewInput(
            task_id="task_001",
            doc_id="doc_001",
            template_id=1,
            template_name="repeat single source",
            template_file="repeat.docx",
            output_file=output_path.name,
            output_path=output_path,
            main_table="customer",
            primary_key_field="customer_id",
            primary_key_value="C001",
            trace_map={
                "customer.id_no": [id_item],
                "customer.address": [address_item],
            },
            final_docx_path=output_path,
        )
    )

    preview_payload = load_json(result.preview_file_path)
    paragraph_runs = [
        run
        for block in preview_payload["blocks"]
        if block["type"] == "paragraph"
        for run in block["runs"]
    ]
    assert [run["trace_id"] for run in paragraph_runs if run["text"] == "ID-001"] == [
        "trace_id_no",
        "trace_id_no",
    ]
    assert [run["trace_id"] for run in paragraph_runs if run["text"] == "ADDR-001"] == [
        "trace_address",
        "trace_address",
    ]

    table_block = next(block for block in preview_payload["blocks"] if block["type"] == "table")
    assert table_block["rows"][0][0]["trace_id"] == "trace_id_no"
    assert table_block["rows"][0][1]["trace_id"] == "trace_address"


def test_field_trace_detail_source_record_highlights_clicked_field(tmp_path) -> None:
    item = make_trace_item("trace_customer_name", "customer_info.customer_name", "customer_info", "customer_name", "Acme")

    from app.engine.trace_builder import build_source_record_view

    record = build_source_record_view([item], item, highlighted_fields={"customer_name"}, highlight_reason="clicked")

    assert record.table_name == "customer_info"
    assert record.excel_row_number == 2
    assert record.fields[0].field_name == "customer_name"
    assert record.fields[0].is_highlighted is True
    assert record.fields[0].highlight_reason == "clicked"


def test_condition_trace_and_preview_block_trace(tmp_path) -> None:
    output_path = tmp_path / "condition.docx"
    template_path = tmp_path / "template.docx"
    write_condition_docx(output_path)
    write_condition_template(template_path)
    house = make_trace_item(
        "trace_house",
        "loan_summary.house_eval_amount",
        "loan_summary",
        "house_eval_amount",
        "900",
    )
    house.raw_value = 900
    balance = make_trace_item("trace_balance", "loan_summary.loan_balance", "loan_summary", "loan_balance", "1250")
    balance.raw_value = 1250

    result = build_trace_and_preview(
        BuildTracePreviewInput(
            task_id="task_001",
            doc_id="doc_001",
            template_id=1,
            template_name="template",
            template_file="template.docx",
            output_file=output_path.name,
            output_path=output_path,
            main_table="loan_summary",
            primary_key_field="branch_name",
            primary_key_value="A支行",
            trace_map={
                "loan_summary.house_eval_amount": [house],
                "loan_summary.loan_balance": [balance],
            },
            final_docx_path=output_path,
            template_path=template_path,
        )
    )

    trace_payload = load_json(result.trace_file_path)
    preview_payload = load_json(result.preview_file_path)

    condition = trace_payload["condition_traces"][0]
    assert condition["expression"] == "loan_summary.house_eval_amount < loan_summary.loan_balance"
    assert condition["used_variables"] == ["loan_summary.house_eval_amount", "loan_summary.loan_balance"]
    assert condition["evaluated_result"] is True
    assert condition["selected_branch"] == "if"
    assert condition["expected_output_text"] == "无法清偿"
    assert condition["is_consistent"] is True
    assert condition["source_records"][0]["fields"][0]["is_highlighted"] is True

    paragraph = preview_payload["blocks"][0]
    assert paragraph["block_trace_id"] == condition["trace_id"]
    assert paragraph["block_trace_kind"] == "condition"


def test_loop_trace_and_preview_table_block_trace(tmp_path) -> None:
    output_path = tmp_path / "loop.docx"
    template_path = tmp_path / "template.docx"
    write_loop_docx(output_path)
    write_loop_template(template_path)
    first_name = make_trace_item("trace_product_a", "products.product_name", "products", "product_name", "产品A", 0)
    first_rate = make_trace_item("trace_rate_a", "products.overdue_rate", "products", "overdue_rate", "1.20%", 0)
    second_name = make_trace_item("trace_product_b", "products.product_name", "products", "product_name", "产品B", 1)
    second_rate = make_trace_item("trace_rate_b", "products.overdue_rate", "products", "overdue_rate", "2.40%", 1)

    result = build_trace_and_preview(
        BuildTracePreviewInput(
            task_id="task_001",
            doc_id="doc_001",
            template_id=1,
            template_name="template",
            template_file="template.docx",
            output_file=output_path.name,
            output_path=output_path,
            main_table="branch_main",
            primary_key_field="branch_name",
            primary_key_value="A支行",
            trace_map={
                "products.product_name": [first_name, second_name],
                "products.overdue_rate": [first_rate, second_rate],
            },
            final_docx_path=output_path,
            template_path=template_path,
        )
    )

    trace_payload = load_json(result.trace_file_path)
    preview_payload = load_json(result.preview_file_path)

    loop = trace_payload["loop_traces"][0]
    assert loop["table_name"] == "products"
    assert loop["loop_alias"] == "p"
    assert loop["used_fields"] == ["products.overdue_rate", "products.product_name"]
    assert loop["matched_row_count"] == 2
    assert len(loop["source_records"]) == 2
    assert any(field["is_highlighted"] for field in loop["source_records"][0]["fields"])

    table_block = preview_payload["blocks"][0]
    assert table_block["block_trace_id"] == loop["trace_id"]
    assert table_block["block_trace_kind"] == "loop"
    assert table_block["rows"][0][0]["trace_id"] == "trace_product_a"
