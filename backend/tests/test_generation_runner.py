from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path
import json
import sqlite3
import sys

from docx import Document
import pandas as pd

PROJECT_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(PROJECT_ROOT))

from app.core.config import AppSettings
from app.engine.generation_runner import GenerateTaskInput, GenerationDependencies, generate_task
from app.models.enums import DataType, RelationType, TableRole, TaskStatus, ValidationStatus
from app.models.validation_models import ValidationReport, ValidationSummary

NOW = datetime(2026, 6, 14, 12, 0, tzinfo=timezone.utc)


def make_settings(tmp_path: Path) -> AppSettings:
    return AppSettings(
        project_root=tmp_path,
        database_path=tmp_path / "config" / "db.sqlite",
        config_dir=tmp_path / "config",
        templates_dir=tmp_path / "templates",
        tasks_dir=tmp_path / "tasks",
        temp_dir=tmp_path / "temp",
    )


def init_connection(tmp_path: Path) -> sqlite3.Connection:
    connection = sqlite3.connect(tmp_path / "config" / "db.sqlite")
    connection.row_factory = sqlite3.Row
    schema_path = PROJECT_ROOT / "app" / "db" / "schema.sql"
    connection.executescript(schema_path.read_text(encoding="utf-8"))
    return connection


def seed_task(connection: sqlite3.Connection, settings: AppSettings, task_id: str = "task_001") -> None:
    now = NOW.isoformat()
    connection.execute(
        """
        INSERT INTO templates (id, template_name, template_file, template_path, main_table, created_at, updated_at)
        VALUES (1, 'due diligence', 'due_diligence.docx', 'templates/due_diligence.docx', 'customer_info', ?, ?)
        """,
        (now, now),
    )
    connection.executemany(
        """
        INSERT INTO template_tables (
            template_id, table_name, table_name_cn, role, relation_type,
            main_join_key, table_join_key, required, sort_order, created_at, updated_at
        )
        VALUES (1, ?, '', ?, ?, ?, ?, 1, ?, ?, ?)
        """,
        [
            ("customer_info", "main", "main", "", "", 0, now, now),
            ("loan_summary", "aux", "one_to_one", "customer_id", "customer_id", 1, now, now),
        ],
    )
    connection.executemany(
        """
        INSERT INTO fields (
            table_name, field_name, data_type, is_primary_key, required, created_at, updated_at
        )
        VALUES (?, ?, ?, ?, 1, ?, ?)
        """,
        [
            ("customer_info", "customer_id", "string", 1, now, now),
            ("customer_info", "customer_name", "string", 0, now, now),
            ("loan_summary", "customer_id", "string", 0, now, now),
            ("loan_summary", "loan_balance", "amount", 0, now, now),
        ],
    )
    task_dir = f"tasks/{task_id}"
    connection.execute(
        """
        INSERT INTO tasks (
            task_id, task_name, template_id, template_name, status, ai_enabled,
            main_table, primary_key_field, task_dir, created_at, updated_at
        )
        VALUES (?, 'task', 1, 'due diligence', 'uploaded', 0, 'customer_info', 'customer_id', ?, ?, ?)
        """,
        (task_id, task_dir, now, now),
    )
    write_task_meta(settings, task_id)


def write_task_meta(settings: AppSettings, task_id: str) -> None:
    task_dir = settings.tasks_dir / task_id
    for directory in (task_dir / "data", task_dir / "output", task_dir / "logs", task_dir / "validation", settings.temp_dir / task_id):
        directory.mkdir(parents=True, exist_ok=True)
    meta = {
        "schema_version": "1.0",
        "task_id": task_id,
        "task_name": "task",
        "template_id": 1,
        "template_name": "due diligence",
        "template_file": "due_diligence.docx",
        "ai_enabled": False,
        "status": "uploaded",
        "main_table": "customer_info",
        "main_table_cn": "",
        "primary_key_field": "customer_id",
        "required_tables": [
            {"table_name": "customer_info", "table_name_cn": "", "role": "main", "relation_type": "main", "main_join_key": "", "table_join_key": "", "required": True},
            {"table_name": "loan_summary", "table_name_cn": "", "role": "aux", "relation_type": "one_to_one", "main_join_key": "customer_id", "table_join_key": "customer_id", "required": True},
        ],
        "uploaded_files": [],
        "output_summary": {"total_rows": 0, "success_count": 0, "failed_count": 0, "warning_count": 0, "error_count": 0, "document_ids": []},
        "paths": {
            "task_dir": f"tasks/{task_id}",
            "data_dir": f"tasks/{task_id}/data",
            "output_dir": f"tasks/{task_id}/output",
            "validation_dir": f"tasks/{task_id}/validation",
            "logs_dir": f"tasks/{task_id}/logs",
        },
        "created_at": NOW.isoformat(),
        "updated_at": NOW.isoformat(),
        "started_at": None,
        "completed_at": None,
    }
    (task_dir / "meta.json").write_text(json.dumps(meta, ensure_ascii=False), encoding="utf-8")


def write_template(settings: AppSettings) -> None:
    settings.templates_dir.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.add_paragraph("Customer: {{ customer_info.customer_name }}")
    document.add_paragraph("Balance: {{ loan_summary.loan_balance }}")
    document.save(settings.templates_dir / "due_diligence.docx")


def write_chinese_template(settings: AppSettings) -> None:
    settings.templates_dir.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.add_paragraph("客户：{{ 客户信息表.客户名称 }}")
    document.add_paragraph("余额：{{ 贷款汇总表.贷款余额 }}")
    document.save(settings.templates_dir / "due_diligence.docx")


def write_excel_data(settings: AppSettings, task_id: str = "task_001") -> None:
    data_dir = settings.tasks_dir / task_id / "data"
    data_dir.mkdir(parents=True, exist_ok=True)
    pd.DataFrame(
        [
            {"customer_id": "C001", "customer_name": "Acme"},
            {"customer_id": "C002", "customer_name": "Beta"},
            {"customer_id": "C003", "customer_name": "Gamma"},
        ]
    ).to_excel(data_dir / "customer_info.xlsx", index=False)
    pd.DataFrame(
        [
            {"customer_id": "C001", "loan_balance": 1200},
            {"customer_id": "C002", "loan_balance": 800},
            {"customer_id": "C003", "loan_balance": 1500},
        ]
    ).to_excel(data_dir / "loan_summary.xlsx", index=False)


def passed_validation_report(task_id: str, *args, **kwargs) -> ValidationReport:
    return ValidationReport(
        task_id=task_id,
        status=ValidationStatus.PASSED,
        summary=ValidationSummary(),
        items=[],
        created_at=NOW,
    )


def setup_project(tmp_path: Path) -> tuple[sqlite3.Connection, AppSettings]:
    settings = make_settings(tmp_path)
    for directory in (settings.config_dir, settings.templates_dir, settings.tasks_dir, settings.temp_dir):
        directory.mkdir(parents=True, exist_ok=True)
    connection = init_connection(tmp_path)
    seed_task(connection, settings)
    write_template(settings)
    write_excel_data(settings)
    connection.commit()
    return connection, settings


def test_generate_task_creates_docx_trace_preview_for_each_main_row(tmp_path) -> None:
    connection, settings = setup_project(tmp_path)

    result = generate_task(
        GenerateTaskInput(task_id="task_001", ai_enabled=False),
        connection=connection,
        settings=settings,
        dependencies=GenerationDependencies(validate_task=passed_validation_report),
    )

    assert result.status == TaskStatus.COMPLETED
    assert result.total_rows == 3
    assert result.success_count == 3
    assert result.failed_count == 0
    output_dir = settings.tasks_dir / "task_001" / "output"
    assert len(list(output_dir.glob("*.docx"))) == 3
    assert len(list(output_dir.glob("*.trace.json"))) == 3
    assert len(list(output_dir.glob("*.preview.json"))) == 3

    documents = connection.execute("SELECT * FROM documents ORDER BY primary_key_value").fetchall()
    assert len(documents) == 3
    assert {row["status"] for row in documents} == {"success"}
    task = connection.execute("SELECT status, total_rows, success_count, failed_count FROM tasks WHERE task_id = 'task_001'").fetchone()
    assert dict(task) == {"status": "completed", "total_rows": 3, "success_count": 3, "failed_count": 0}


def test_generate_task_renders_chinese_template_variables_and_keeps_original_trace_path(tmp_path) -> None:
    connection, settings = setup_project(tmp_path)
    now = NOW.isoformat()
    connection.execute("UPDATE template_tables SET table_name_cn = '客户信息表' WHERE table_name = 'customer_info'")
    connection.execute("UPDATE template_tables SET table_name_cn = '贷款汇总表' WHERE table_name = 'loan_summary'")
    connection.execute(
        "UPDATE fields SET table_name_cn = '客户信息表', field_name_cn = '客户编号', updated_at = ? WHERE table_name = 'customer_info' AND field_name = 'customer_id'",
        (now,),
    )
    connection.execute(
        "UPDATE fields SET table_name_cn = '客户信息表', field_name_cn = '客户名称', updated_at = ? WHERE table_name = 'customer_info' AND field_name = 'customer_name'",
        (now,),
    )
    connection.execute(
        "UPDATE fields SET table_name_cn = '贷款汇总表', field_name_cn = '客户编号', updated_at = ? WHERE table_name = 'loan_summary' AND field_name = 'customer_id'",
        (now,),
    )
    connection.execute(
        "UPDATE fields SET table_name_cn = '贷款汇总表', field_name_cn = '贷款余额', updated_at = ? WHERE table_name = 'loan_summary' AND field_name = 'loan_balance'",
        (now,),
    )
    connection.commit()
    write_chinese_template(settings)

    result = generate_task(
        GenerateTaskInput(task_id="task_001", ai_enabled=False),
        connection=connection,
        settings=settings,
        dependencies=GenerationDependencies(validate_task=passed_validation_report),
    )

    assert result.status == TaskStatus.COMPLETED
    assert result.success_count == 3
    output_dir = settings.tasks_dir / "task_001" / "output"
    first_docx = sorted(output_dir.glob("*.docx"))[0]
    paragraphs = [paragraph.text for paragraph in Document(first_docx).paragraphs]
    assert "客户：Acme" in paragraphs

    first_trace = json.loads(sorted(output_dir.glob("*.trace.json"))[0].read_text(encoding="utf-8"))
    customer_name_trace = next(
        item
        for item in first_trace["trace_items"]
        if item["canonical_var_path"] == "customer_info.customer_name"
    )
    assert customer_name_trace["var_path"] == "customer_info.customer_name"
    assert customer_name_trace["original_var_path"] == "客户信息表.客户名称"


def test_single_document_failure_does_not_stop_other_rows(tmp_path) -> None:
    connection, settings = setup_project(tmp_path)

    def flaky_render(template_path, context, output_path):
        from app.engine.docx_renderer import render_docx_template

        if context["customer_info"]["customer_id"] == "C002":
            return type("RenderResult", (), {"success": False, "error_message": "boom"})()
        return render_docx_template(template_path, context, output_path)

    result = generate_task(
        GenerateTaskInput(task_id="task_001", ai_enabled=False),
        connection=connection,
        settings=settings,
        dependencies=GenerationDependencies(
            validate_task=passed_validation_report,
            render_docx_template=flaky_render,
        ),
    )

    assert result.status == TaskStatus.PARTIAL_FAILED
    assert result.success_count == 2
    assert result.failed_count == 1
    output_dir = settings.tasks_dir / "task_001" / "output"
    assert len(list(output_dir.glob("*.docx"))) == 2
    assert len(list(output_dir.glob("*.trace.json"))) == 2
    assert len(list(output_dir.glob("*.preview.json"))) == 2

    documents = connection.execute("SELECT primary_key_value, status FROM documents ORDER BY primary_key_value").fetchall()
    assert [(row["primary_key_value"], row["status"]) for row in documents] == [
        ("C001", "success"),
        ("C002", "failed"),
        ("C003", "success"),
    ]
    task_status = connection.execute("SELECT status FROM tasks WHERE task_id = 'task_001'").fetchone()["status"]
    assert task_status == "partial_failed"


def test_all_document_failures_mark_task_failed(tmp_path) -> None:
    connection, settings = setup_project(tmp_path)

    def failing_render(template_path, context, output_path):
        return type("RenderResult", (), {"success": False, "error_message": "boom"})()

    result = generate_task(
        GenerateTaskInput(task_id="task_001", ai_enabled=False),
        connection=connection,
        settings=settings,
        dependencies=GenerationDependencies(
            validate_task=passed_validation_report,
            render_docx_template=failing_render,
        ),
    )

    assert result.status == TaskStatus.FAILED
    assert result.success_count == 0
    assert result.failed_count == 3
    task_status = connection.execute("SELECT status FROM tasks WHERE task_id = 'task_001'").fetchone()["status"]
    assert task_status == "failed"
