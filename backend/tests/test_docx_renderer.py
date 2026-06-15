from pathlib import Path
import sys

from docx import Document

PROJECT_ROOT = Path(__file__).resolve().parents[1]
sys.path.insert(0, str(PROJECT_ROOT))

from app.engine.docx_renderer import RenderDocxInput, render_docx, render_docx_template


def paragraph_text(document: Document) -> list[str]:
    return [paragraph.text for paragraph in document.paragraphs]


def table_rows(document: Document) -> list[list[str]]:
    rows: list[list[str]] = []
    for table in document.tables:
        for row in table.rows:
            rows.append([cell.text for cell in row.cells])
    return rows


def write_basic_template(path: Path) -> None:
    document = Document()
    document.add_paragraph("客户名称：{{ customer_info.customer_name }}")
    document.add_paragraph(
        "{% if loan_summary.bad_rate > 0.02 %}风险偏高{% else %}风险可控{% endif %}"
    )
    document.save(path)


def write_table_loop_template(path: Path) -> None:
    document = Document()
    table = document.add_table(rows=3, cols=2)
    table.rows[0].cells[0].text = "{%tr for item in collateral_info %}"
    table.rows[1].cells[0].text = "{{ item.collateral_name }}"
    table.rows[1].cells[1].text = "{{ item.eval_amount }}"
    table.rows[2].cells[0].text = "{%tr endfor %}"
    document.save(path)


def write_merged_table_template(path: Path) -> None:
    document = Document()
    table = document.add_table(rows=3, cols=3)
    table.cell(0, 0).merge(table.cell(0, 1)).text = "Merged {{ customer_info.customer_name }}"
    table.cell(0, 2).text = "Header"
    table.cell(1, 0).merge(table.cell(2, 0)).text = "Vertical"
    table.cell(1, 1).text = "{{ loan_summary.loan_balance }}"
    table.cell(1, 2).text = "Tail"
    table.cell(2, 1).text = "Bottom"
    table.cell(2, 2).text = "Right"
    document.save(path)


def has_grid_span(document: Document) -> bool:
    return bool(document.element.xpath(".//w:gridSpan"))


def has_vertical_merge(document: Document) -> bool:
    return bool(document.element.xpath(".//w:vMerge"))


def test_render_docx_template_replaces_variables_and_if_conditions(tmp_path) -> None:
    template_path = tmp_path / "template.docx"
    output_path = tmp_path / "output.docx"
    write_basic_template(template_path)

    result = render_docx_template(
        template_path,
        {
            "customer_info": {"customer_name": "Acme"},
            "loan_summary": {"bad_rate": 0.018},
        },
        output_path,
    )

    assert result.success is True
    assert result.error_message == ""
    rendered = Document(output_path)
    texts = paragraph_text(rendered)
    assert "客户名称：Acme" in texts
    assert "风险可控" in texts


def test_render_docx_template_supports_true_if_branch(tmp_path) -> None:
    template_path = tmp_path / "template.docx"
    output_path = tmp_path / "output.docx"
    write_basic_template(template_path)

    result = render_docx(
        RenderDocxInput(
            template_path=template_path,
            context={
                "customer_info": {"customer_name": "Beta"},
                "loan_summary": {"bad_rate": 0.032},
            },
            output_path=output_path,
        )
    )

    assert result.success is True
    assert "风险偏高" in paragraph_text(Document(output_path))


def test_render_docx_template_supports_docxtpl_table_row_loop(tmp_path) -> None:
    template_path = tmp_path / "table_template.docx"
    output_path = tmp_path / "table_output.docx"
    write_table_loop_template(template_path)

    result = render_docx_template(
        template_path,
        {
            "collateral_info": [
                {"collateral_name": "Factory", "eval_amount": "3000.00"},
                {"collateral_name": "Equipment", "eval_amount": "500.00"},
            ]
        },
        output_path,
    )

    assert result.success is True
    rendered = Document(output_path)
    assert table_rows(rendered) == [["Factory", "3000.00"], ["Equipment", "500.00"]]


def test_render_docx_template_preserves_merged_table_structure(tmp_path) -> None:
    template_path = tmp_path / "merged_template.docx"
    output_path = tmp_path / "merged_output.docx"
    write_merged_table_template(template_path)

    result = render_docx_template(
        template_path,
        {
            "customer_info": {"customer_name": "Acme"},
            "loan_summary": {"loan_balance": "3000.00"},
        },
        output_path,
    )

    assert result.success is True
    rendered = Document(output_path)
    assert rendered.tables[0].cell(0, 0).text == "Merged Acme"
    assert has_grid_span(rendered)
    assert has_vertical_merge(rendered)


def test_rendered_docx_can_be_opened_by_python_docx(tmp_path) -> None:
    template_path = tmp_path / "template.docx"
    output_path = tmp_path / "nested" / "output.docx"
    write_basic_template(template_path)

    result = render_docx_template(
        template_path,
        {
            "customer_info": {"customer_name": "Acme"},
            "loan_summary": {"bad_rate": 0.018},
        },
        output_path,
    )

    assert result.success is True
    opened = Document(output_path)
    assert opened.paragraphs[0].text == "客户名称：Acme"


def test_render_docx_template_returns_clear_error_without_output_for_invalid_template(tmp_path) -> None:
    template_path = tmp_path / "invalid_template.docx"
    output_path = tmp_path / "output.docx"
    document = Document()
    document.add_paragraph("{% if customer_info.customer_name %}missing endif")
    document.save(template_path)

    result = render_docx_template(
        template_path,
        {"customer_info": {"customer_name": "Acme"}},
        output_path,
    )

    assert result.success is False
    assert "Failed to render docx template" in result.error_message
    assert "invalid_template.docx" in result.error_message
    assert not output_path.exists()


def test_render_docx_template_reports_missing_template_file(tmp_path) -> None:
    output_path = tmp_path / "output.docx"

    result = render_docx_template(
        tmp_path / "missing.docx",
        {},
        output_path,
    )

    assert result.success is False
    assert result.error_message == f"Template file not found: {tmp_path / 'missing.docx'}"
    assert not output_path.exists()
